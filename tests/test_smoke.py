"""اختبارات دخان لمقيّم خطة البحث — تعمل دون مفاتيح API ودون شبكة."""
import io
import json
import os
import sys
import unittest
from unittest import mock

sys.path.insert(
    0,
    os.path.dirname(
        os.path.dirname(os.path.abspath(__file__))
    ),
)

from criteria import CRITERIA, CATEGORIES, MAX_SCORE  # noqa: E402
import app  # noqa: E402


class TestCore(unittest.TestCase):
    def test_criteria_counts(self):
        self.assertEqual(len(CRITERIA), 44)
        self.assertEqual(MAX_SCORE, 88)
        ids = [item[0] for item in CRITERIA]
        self.assertIn(12, ids)
        self.assertIn(51, ids)
        self.assertIn(68, ids)
        self.assertIn(71, ids)
        self.assertEqual(len(CATEGORIES), 11)

    def test_build_prompt_contains_all_ids(self):
        prompt = app.build_prompt("نص تجريبي")
        for number, _, _ in CRITERIA:
            self.assertIn(f"{number}.", prompt)

    def test_parse_json_fenced_and_plain(self):
        data = app.parse_json('```json\n{"results": []}\n```')
        self.assertEqual(data, {"results": []})
        data = app.parse_json(
            'نص قبلي {"results": [], "x": 1} نص بعدي'
        )
        self.assertEqual(data, {"results": [], "x": 1})

    def test_normalize_fills_missing(self):
        payload = {
            "overall_analysis": "تحليل عام",
            "results": [
                {
                    "id": 12,
                    "score": 2,
                    "status": "متحقق",
                    "evidence": "دليل",
                    "explanation": "شرح",
                    "suggestion": "توصية",
                }
            ],
        }
        results, analysis = app.normalize_results(payload)
        self.assertEqual(len(results), 44)
        self.assertEqual(analysis, "تحليل عام")
        self.assertEqual(results[0]["score"], 2)
        missing = [r for r in results if r["id"] != 12]
        self.assertTrue(all(r["score"] == 0 for r in missing))

    def test_calculate_levels(self):
        total, maximum, percentage, level = app.calculate(
            [{"score": 2}] * 44
        )
        self.assertEqual((total, maximum), (88, 88))
        self.assertEqual(level, "ممتاز")

    def test_openai_honors_custom_base_url(self):
        payload = {
            "choices": [
                {
                    "message": {
                        "content": json.dumps(
                            {
                                "overall_analysis": "أ",
                                "results": [],
                            }
                        )
                    }
                }
            ]
        }

        def fake_secret(name, default=""):
            if name == "OPENAI_BASE_URL":
                return "https://api.example.com/v1"
            return default

        with mock.patch("app.get_secret", side_effect=fake_secret), \
                mock.patch("requests.post") as post:
            post.return_value.status_code = 200
            post.return_value.json.return_value = payload
            results, _ = app.call_openai("نص", "KEY", "gpt-test")
            self.assertEqual(
                post.call_args[0][0],
                "https://api.example.com/v1/chat/completions",
            )
            self.assertEqual(len(results), 44)

    def test_anthropic_sends_temperature(self):
        with mock.patch("requests.post") as post:
            response = mock.MagicMock()
            response.status_code = 200
            response.json.return_value = {
                "content": [
                    {
                        "type": "text",
                        "text": json.dumps(
                            {
                                "overall_analysis": "أ",
                                "results": [],
                            }
                        ),
                    }
                ]
            }
            post.return_value = response
            app.call_anthropic("نص", "KEY", "claude-opus-4-8")
            body = post.call_args.kwargs["json"]
            self.assertEqual(body["temperature"], 0.1)
            self.assertEqual(body["max_tokens"], 24000)


class TestAppFlow(unittest.TestCase):
    def test_end_to_end_basic_mode(self):
        from streamlit.testing.v1 import AppTest
        from docx import Document

        document = Document()
        document.add_paragraph("عنوان البحث: دراسة تجريبية")
        document.add_paragraph(
            "المقدمة: " + "نص المقدمة الأكاديمية لخطة البحث. " * 15
        )
        document.add_paragraph(
            "مشكلة الدراسة: " + "وصف مشكلة البحث ومبرراتها. " * 15
        )
        document.add_paragraph(
            "أهداف الدراسة: " + "أهداف البحث العلمي المحددة. " * 15
        )
        document.add_paragraph(
            "منهج البحث: " + "المنهج الوصفي التحليلي المناسب. " * 15
        )
        buffer = io.BytesIO()
        document.save(buffer)
        docx_bytes = buffer.getvalue()

        app_path = os.path.join(
            os.path.dirname(
                os.path.dirname(os.path.abspath(__file__))
            ),
            "app.py",
        )

        at = AppTest.from_file(app_path, default_timeout=180)
        at.run()
        self.assertEqual(len(at.exception), 0)

        at.radio[0].set_value("🔎 فحص أولي")
        at.run()

        at.file_uploader[0].set_value(
            [
                (
                    "خطة.docx",
                    docx_bytes,
                    "application/vnd.openxmlformats-"
                    "officedocument.wordprocessingml.document",
                )
            ]
        )
        at.run()
        self.assertEqual(len(at.exception), 0)

        at.button[0].click()
        at.run()

        self.assertEqual(
            len(at.exception),
            0,
            [repr(entry.value) for entry in at.exception],
        )
        self.assertGreaterEqual(len(at.metric), 4)


if __name__ == "__main__":
    unittest.main(verbosity=2)
