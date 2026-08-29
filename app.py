import io
import json
import os
import re
from datetime import datetime

import streamlit as st

# ============================================================

# المعايير المعتمدة حصراً

# النطاق: 12–51 و68–71 فقط

# ============================================================

CRITERIA = [
(12, "صياغة العنوان", "يجب أن يكون مختصر وواضح ومحدد."),
(13, "صياغة العنوان", "يتضح منه الموضوع الرئيسي للدراسة."),
(14, "صياغة العنوان", "يحتوي على المتغيرات الرئيسية للدراسة."),
(15, "صياغة العنوان", "لا يحتوي على عبارات تحتمل أكثر من معنى."),

(16, "كتابة المقدمة", "مكتوبة بأسلوب منطقي من العام للخاص."),
(17, "كتابة المقدمة", "توضح فكرة البحث والموضوع والعوامل المرتبطة به."),
(18, "كتابة المقدمة", "خالية من الإسهاب الممل والإيجاز المخل."),
(19, "كتابة المقدمة", "تتضح من خلالها شخصية الباحث."),

(20, "صياغة مشكلة الدراسة", "الكتابة بأسلوب يقنع القارئ بوجود مشكلة فعلاً."),
(21, "صياغة مشكلة الدراسة", "يتضح من خلالها مبررات اختيار المشكلة."),
(22, "صياغة مشكلة الدراسة", "يتضح من خلالها مصادر الوصول للمشكلة."),
(23, "صياغة مشكلة الدراسة", "تنتهي بسؤال بحثي أو عبارة تقريرية تحدد مشكلة الدراسة."),

(24, "صياغة تساؤلات الدراسة", "تتفرع من سؤال رئيس محدد لتغطي جميع جوانب البحث."),
(25, "صياغة تساؤلات الدراسة", "تمثل تساؤلات بحثية فعلاً."),
(26, "صياغة تساؤلات الدراسة", "متسلسلة حسب مستويات أهداف العلم."),
(27, "صياغة تساؤلات الدراسة", "واضحة ومحددة وغير مركبة."),

(28, "صياغة فروض الدراسة", "واضحة ومحددة وقابلة للاختبار."),
(29, "صياغة فروض الدراسة", "ذات علاقة بمتغيرات الدراسة."),
(30, "صياغة فروض الدراسة", "متنوعة بين الفروض الصفرية والبديلة."),
(31, "صياغة فروض الدراسة", "تمثل تخميناً ذكياً وبعيدة عن العشوائية."),

(32, "صياغة أهمية الدراسة", "مكتوبة وتوضح الأهمية العلمية للبحث والدراسة."),
(33, "صياغة أهمية الدراسة", "مكتوبة في عبارات محددة وواضحة بعيدة عن الإسهاب والتشتت."),
(34, "صياغة أهمية الدراسة", "تتضح من خلالها الأهمية النظرية للدراسة."),
(35, "صياغة أهمية الدراسة", "تتضح من خلالها الأهمية التطبيقية للدراسة."),

(36, "صياغة أهداف الدراسة", "واضحة ومحددة وغير مركبة."),
(37, "صياغة أهداف الدراسة", "ذات علاقة مباشرة بتساؤلات الدراسة وفروضها."),
(38, "صياغة أهداف الدراسة", "مرتبة حسب مستويات أهداف العلم."),
(39, "صياغة أهداف الدراسة", "تحقق في مجملها الهدف العام للدراسة."),

(40, "كتابة مصطلحات الدراسة", "يجب أن تغطي جميع المصطلحات المتعلقة بموضوع الدراسة وعنوانها."),
(41, "كتابة مصطلحات الدراسة", "تتضمن التعريفات اللغوية والاصطلاحية والإجرائية."),
(42, "كتابة مصطلحات الدراسة", "يجب أن تكون التعريفات محددة وموجزة ومكتوبة بشكل سهل."),
(43, "كتابة مصطلحات الدراسة", "مرتبة بشكل منطقي وسليم."),

(44, "صياغة حدود الدراسة", "نطاق الدراسة محدد بشكل دقيق."),
(45, "صياغة حدود الدراسة", "تشمل جميع الأبعاد التي تحتاج إلى تحديد."),
(46, "صياغة حدود الدراسة", "مكتوبة بشكل واضح ودقيق دون أي لبس."),
(47, "صياغة حدود الدراسة", "تتضح من خلالها إمكانية تعميم نتائج الدراسة من عدمها."),

(48, "اختيار منهج البحث", "اختيار المنهج الملائم لطبيعة البحث."),
(49, "اختيار منهج البحث", "مبررات اختياره مذكورة ومحددة."),
(50, "اختيار منهج البحث", "خطوات تنفيذه متسلسلة بشكل منطقي."),
(51, "اختيار منهج البحث", "يتم توظيفه خلال الدراسة بشكل ملائم."),

(68, "أخرى", "الخطة مترابطة وتمثل وحدة بحثية متناسقة."),
(69, "أخرى", "تنسيق الخطة سليم وملائم."),
(70, "أخرى", "شخصية الباحث واضحة في محاور الخطة."),
(71, "أخرى", "لغة الكتابة سليمة إملائياً ونحوياً."),
```

]

CATEGORIES = {}

for item in CRITERIA:
CATEGORIES.setdefault(
item[1],
[]
).append(item)

MAX_SCORE = len(CRITERIA) * 2

# ============================================================

# إعداد الصفحة

# ============================================================

st.set_page_config(
page_title="مقيّم خطة البحث العلمي",
page_icon="📚",
layout="wide",
initial_sidebar_state="expanded",
)

# ============================================================

# CSS

# ============================================================

st.markdown(
""" <style>

```
.block-container {
    direction: rtl;
    max-width: 1450px;
    padding-top: 1rem;
}

.hero {
    padding: 30px;
    border-radius: 20px;
    background: linear-gradient(
        135deg,
        #173b63,
        #2f7f8f
    );
    color: white;
    margin-bottom: 25px;
}

.hero h1 {
    margin: 0;
    font-size: 2.2rem;
}

.hero p {
    margin-top: 12px;
    line-height: 1.9;
    font-size: 1.05rem;
}

.result-card {
    padding: 20px;
    border-radius: 16px;
    border: 1px solid #dce5e9;
    background: white;
    margin-bottom: 15px;
    line-height: 1.9;
}

</style>
""",
unsafe_allow_html=True,
```

)

# ============================================================

# العنوان

# ============================================================

st.markdown(
""" <div class="hero"> <h1>📚 مقيّم خطة البحث العلمي</h1>

```
    <p>
    تقييم أكاديمي لخطة البحث باستخدام المعايير
    من 12 إلى 51 ومن 68 إلى 71 فقط،
    مع استخراج النص العربي من الملفات،
    وإعطاء درجة رقمية وتحليل أكاديمي مكتوب.
    </p>
</div>
""",
unsafe_allow_html=True,
```

)

# ============================================================

# Secrets

# ============================================================

def get_secret(
name,
default=""
):

```
try:

    if name in st.secrets:

        return str(
            st.secrets[name]
        )

except Exception:

    pass

return os.getenv(
    name,
    default
)
```

# ============================================================

# تنظيف النص

# ============================================================

def clean_text(text):

```
if not text:

    return ""

text = text.replace(
    "\u200f",
    " "
)

text = text.replace(
    "\u200e",
    " "
)

text = text.replace(
    "\ufeff",
    " "
)

text = re.sub(
    r"[ \t]+",
    " ",
    text
)

text = re.sub(
    r"\n{3,}",
    "\n\n",
    text
)

return text.strip()
```

# ============================================================

# استخراج Word

# ============================================================

@st.cache_data(
show_spinner=False
)
def extract_docx(data):

```
from docx import Document

document = Document(
    io.BytesIO(data)
)

parts = []

for paragraph in document.paragraphs:

    if paragraph.text.strip():

        parts.append(
            paragraph.text.strip()
        )

for table in document.tables:

    for row in table.rows:

        cells = []

        for cell in row.cells:

            value = (
                cell.text
                .strip()
            )

            cells.append(
                value
            )

        if any(cells):

            parts.append(
                " | ".join(cells)
            )

return clean_text(
    "\n".join(parts)
)
```

# ============================================================

# استخراج PDF النصي

# ============================================================

@st.cache_data(
show_spinner=False
)
def extract_pdf_text(data):

```
import fitz

document = fitz.open(
    stream=data,
    filetype="pdf"
)

pages = []

for page in document:

    text = page.get_text(
        "text"
    )

    if text:

        pages.append(
            text
        )

document.close()

return clean_text(
    "\n".join(pages)
)
```

# ============================================================

# تحميل PaddleOCR مرة واحدة

# ============================================================

@st.cache_resource(
show_spinner="جاري تحميل نموذج OCR العربي لأول مرة..."
)
def load_arabic_ocr():

```
from paddleocr import PaddleOCR

ocr = PaddleOCR(
    lang="ar",
    use_doc_orientation_classify=True,
    use_doc_unwarping=True,
    use_textline_orientation=True
)

return ocr
```

# ============================================================

# تحسين الصورة قبل OCR

# ============================================================

def preprocess_image(
image
):

```
import cv2
import numpy as np

image = np.array(
    image
)

if len(image.shape) == 3:

    image = cv2.cvtColor(
        image,
        cv2.COLOR_RGB2BGR
    )

height, width = (
    image.shape[:2]
)

# لا نضخم الصور الضخمة أكثر من اللازم
max_dimension = 3000

current_max = max(
    height,
    width
)

if current_max < 1800:

    scale = 2.0

elif current_max < 3000:

    scale = 1.5

else:

    scale = 1.0

if scale != 1.0:

    image = cv2.resize(
        image,
        None,
        fx=scale,
        fy=scale,
        interpolation=cv2.INTER_CUBIC
    )

# تحسين التباين
lab = cv2.cvtColor(
    image,
    cv2.COLOR_BGR2LAB
)

l_channel, a_channel, b_channel = (
    cv2.split(lab)
)

clahe = cv2.createCLAHE(
    clipLimit=2.0,
    tileGridSize=(8, 8)
)

l_channel = clahe.apply(
    l_channel
)

image = cv2.merge(
    (
        l_channel,
        a_channel,
        b_channel
    )
)

image = cv2.cvtColor(
    image,
    cv2.COLOR_LAB2BGR
)

return image
```

# ============================================================

# استخراج النص من نتيجة PaddleOCR

# ============================================================

def extract_paddle_result_text(
result
):

```
lines = []

for res in result:

    try:

        # PaddleOCR 3.x
        if hasattr(
            res,
            "json"
        ):

            data = res.json

            if callable(data):

                data = data()

        else:

            data = res

    except Exception:

        data = {}

    if not isinstance(
        data,
        dict
    ):

        continue

    if "res" in data:

        data = data["res"]

    rec_texts = data.get(
        "rec_texts",
        []
    )

    rec_scores = data.get(
        "rec_scores",
        []
    )

    for index, text in enumerate(
        rec_texts
    ):

        text = str(
            text
        ).strip()

        if not text:

            continue

        score = 1.0

        if index < len(
            rec_scores
        ):

            try:

                score = float(
                    rec_scores[index]
                )

            except Exception:

                score = 1.0

        # الاحتفاظ بالنتائج المقبولة
        if score >= 0.35:

            lines.append(
                text
            )

return lines
```

# ============================================================

# OCR صورة

# ============================================================

def extract_image_paddle(
data
):

```
from PIL import Image

image = Image.open(
    io.BytesIO(data)
).convert("RGB")

image = preprocess_image(
    image
)

ocr = load_arabic_ocr()

result = ocr.predict(
    image
)

lines = (
    extract_paddle_result_text(
        result
    )
)

return clean_text(
    "\n".join(lines)
)
```

# ============================================================

# OCR PDF ممسوح

# ============================================================

def extract_pdf_paddle(
data
):

```
import fitz
from PIL import Image

document = fitz.open(
    stream=data,
    filetype="pdf"
)

total_pages = len(
    document
)

pages = []

ocr = load_arabic_ocr()

progress = st.progress(
    0
)

status = st.empty()

for page_number, page in enumerate(
    document
):

    status.write(
        f"🔎 معالجة الصفحة "
        f"{page_number + 1} "
        f"من {total_pages}..."
    )

    # 2.5x للحصول على صورة جيدة
    pixmap = page.get_pixmap(
        matrix=fitz.Matrix(
            2.5,
            2.5
        ),
        alpha=False
    )

    image = Image.open(
        io.BytesIO(
            pixmap.tobytes(
                "png"
            )
        )
    ).convert("RGB")

    image = preprocess_image(
        image
    )

    result = ocr.predict(
        image
    )

    page_lines = (
        extract_paddle_result_text(
            result
        )
    )

    pages.append(
        "\n".join(
            page_lines
        )
    )

    progress.progress(
        (page_number + 1)
        / total_pages
    )

document.close()

status.empty()
progress.empty()

return clean_text(
    "\n\n".join(
        pages
    )
)
```

# ============================================================

# تحديد نوع الملف

# ============================================================

def extract_file(
uploaded_file
):

```
data = (
    uploaded_file
    .getvalue()
)

filename = (
    uploaded_file.name
    .lower()
)

# Word
if filename.endswith(
    ".docx"
):

    return (
        extract_docx(data),
        "Word — استخراج النص الأصلي"
    )

# PDF
if filename.endswith(
    ".pdf"
):

    text = extract_pdf_text(
        data
    )

    compact = re.sub(
        r"\s",
        "",
        text
    )

    # إذا كان PDF نصياً
    if len(compact) >= 500:

        return (
            text,
            "PDF — استخراج النص الأصلي"
        )

    # إذا كان PDF صورة
    return (
        extract_pdf_paddle(
            data
        ),
        "PDF ممسوح — PaddleOCR عربي"
    )

# صور
if filename.endswith(
    (
        ".png",
        ".jpg",
        ".jpeg",
        ".webp"
    )
):

    return (
        extract_image_paddle(
            data
        ),
        "صورة — PaddleOCR عربي"
    )

raise ValueError(
    "نوع الملف غير مدعوم."
)
```

# ============================================================

# تحديد عنوان البحث

# ============================================================

def find_title(
text
):

```
lines = [
    line.strip()
    for line in text.splitlines()
    if line.strip()
]

patterns = [
    r"^عنوان البحث",
    r"^عنوان الدراسة",
    r"^عنوان",
    r"^title"
]

for index, line in enumerate(
    lines[:100]
):

    for pattern in patterns:

        if re.search(
            pattern,
            line,
            re.IGNORECASE
        ):

            if index + 1 < len(
                lines
            ):

                return lines[
                    index + 1
                ]

if lines:

    return lines[0]

return "غير محدد"
```

# ============================================================

# استخراج المقدمة

# ============================================================

def get_introduction(
text
):

```
lines = [
    line.strip()
    for line in text.splitlines()
    if line.strip()
]

start = None

for index, line in enumerate(
    lines
):

    if any(
        line.startswith(x)
        for x in [
            "المقدمة",
            "مقدمة",
            "تمهيد"
        ]
    ):

        start = index + 1

        break

if start is None:

    return ""

stops = [
    "مشكلة الدراسة",
    "إشكالية الدراسة",
    "تساؤلات الدراسة",
    "أسئلة الدراسة",
    "فرضيات الدراسة",
    "فروض الدراسة",
    "أهمية الدراسة",
    "أهداف الدراسة",
    "مصطلحات الدراسة",
    "حدود الدراسة",
    "منهج الدراسة"
]

result = []

for line in lines[start:]:

    if any(
        line.startswith(stop)
        for stop in stops
    ):

        break

    result.append(
        line
    )

    if len(
        " ".join(result)
    ) >= 9000:

        break

return " ".join(
    result
)[:9000]
```

# ============================================================

# المعايير في Prompt

# ============================================================

def criteria_for_prompt():

```
return "\n".join(
    f"{number}. {category}: {criterion}"
    for number, category, criterion
    in CRITERIA
)
```

# ============================================================

# Prompt التقييم

# ============================================================

def build_prompt(
text
):

```
return f"""
```

أنت أستاذ جامعي ومحكّم متخصص في مناهج البحث العلمي.

مهمتك تقييم خطة البحث المرفوعة.

مهم جداً:
استخدم المعايير التالية فقط.
لا تضف أي معيار آخر.

نطاق التقييم:
المعايير 12 إلى 51
والمعايير 68 إلى 71.

عدد المعايير:
{len(CRITERIA)}

الدرجة لكل معيار:

0 = غير متحقق
1 = متحقق جزئياً
2 = متحقق بوضوح

يجب عدم منح الدرجة لمجرد وجود كلمة مفتاحية.
يجب قراءة المعنى والسياق والترابط.

إذا لم يوجد دليل كافٍ في النص:
اكتب:
"لا يوجد دليل كافٍ في النص."

لكل معيار أعد:

id
score
status
evidence
explanation
suggestion

الحكم:

غير متحقق
متحقق جزئياً
متحقق

evidence:
دليل حقيقي ومحدد من نص الخطة.

explanation:
تحليل أكاديمي يوضح سبب الدرجة.

suggestion:
توصية عملية لتحسين المعيار.

==================================================
التحليل العام
=============

أعد أيضاً:

overall_analysis

ويجب أن يتناول:

1. العنوان
2. المقدمة
3. مشكلة الدراسة
4. تساؤلات الدراسة
5. فروض الدراسة
6. أهمية الدراسة
7. أهداف الدراسة
8. مصطلحات الدراسة
9. حدود الدراسة
10. منهج البحث
11. ترابط الخطة
12. شخصية الباحث
13. اللغة والأسلوب
14. نقاط القوة
15. نقاط الضعف
16. أهم التعديلات المقترحة

==================================================
عنوان البحث
===========

{find_title(text)}

==================================================
المقدمة
=======

{get_introduction(text)}

==================================================
المعايير
========

{criteria_for_prompt()}

==================================================
نص خطة البحث
============

{text[:70000]}

==================================================

أعد JSON فقط دون Markdown.

الصيغة المطلوبة:

{{
"overall_analysis": "...",
"results": [
{{
"id": 12,
"score": 2,
"status": "متحقق",
"evidence": "...",
"explanation": "...",
"suggestion": "..."
}}
]
}}

يجب أن يحتوي results على جميع المعايير
الـ {len(CRITERIA)}.

لا تضف أي معيار غير موجود في القائمة.
"""

# ============================================================

# تنظيف JSON

# ============================================================

def parse_json(
text
):

````
text = text.strip()

text = re.sub(
    r"^```json\s*",
    "",
    text,
    flags=re.IGNORECASE
)

text = re.sub(
    r"^```\s*",
    "",
    text
)

text = re.sub(
    r"\s*```$",
    "",
    text
)

try:

    return json.loads(
        text
    )

except Exception:

    start = text.find(
        "{"
    )

    end = text.rfind(
        "}"
    )

    if (
        start == -1
        or end == -1
    ):

        raise ValueError(
            "لم يتم العثور على JSON صالح."
        )

    return json.loads(
        text[
            start:end + 1
        ]
    )
````

# ============================================================

# توحيد النتائج

# ============================================================

def normalize_results(
data
):

```
allowed = {
    item[0]: item
    for item in CRITERIA
}

received = {}

for item in data.get(
    "results",
    []
):

    try:

        number = int(
            item.get("id")
        )

    except Exception:

        continue

    if number not in allowed:

        continue

    try:

        score = int(
            item.get(
                "score",
                0
            )
        )

    except Exception:

        score = 0

    score = max(
        0,
        min(
            2,
            score
        )
    )

    status_default = [
        "غير متحقق",
        "متحقق جزئياً",
        "متحقق"
    ][score]

    received[number] = {
        "id": number,
        "category": allowed[number][1],
        "criterion": allowed[number][2],
        "score": score,
        "status": item.get(
            "status",
            status_default
        ),
        "evidence": item.get(
            "evidence",
            "لا يوجد دليل كافٍ في النص."
        ),
        "explanation": item.get(
            "explanation",
            ""
        ),
        "suggestion": item.get(
            "suggestion",
            ""
        )
    }

results = []

for (
    number,
    category,
    criterion
) in CRITERIA:

    if number in received:

        results.append(
            received[number]
        )

    else:

        results.append(
            {
                "id": number,
                "category": category,
                "criterion": criterion,
                "score": 0,
                "status": "غير متحقق",
                "evidence":
                    "لم يقدم النموذج دليلاً لهذا المعيار.",
                "explanation":
                    "لم يتم الحصول على نتيجة كافية.",
                "suggestion":
                    "يرجى مراجعة هذا المعيار."
            }
        )

return (
    results,
    data.get(
        "overall_analysis",
        "لم يتم توفير تحليل عام."
    )
)
```

# ============================================================

# KiosAPI

# OpenAI-compatible

# ============================================================

def call_kiosapi(
text,
api_key,
model
):

```
import requests

api_key = (
    api_key
    .strip()
)

if not api_key:

    raise ValueError(
        "مفتاح KiosAPI فارغ."
    )

url = (
    "https://api.kiosapi.id/v1/"
    "chat/completions"
)

response = requests.post(

    url,

    headers={
        "Authorization":
            f"Bearer {api_key}",

        "Content-Type":
            "application/json"
    },

    json={
        "model": model,

        "messages": [
            {
                "role": "system",
                "content":
                    "أنت محكّم أكاديمي "
                    "متخصص في مناهج البحث العلمي. "
                    "أجب بالعربية وأعد JSON فقط."
            },
            {
                "role": "user",
                "content":
                    build_prompt(text)
            }
        ],

        "temperature": 0.1
    },

    timeout=300
)

if response.status_code != 200:

    try:

        error_data = (
            response.json()
        )

        error_message = (
            error_data
            .get(
                "error",
                {}
            )
            .get(
                "message",
                response.text
            )
        )

    except Exception:

        error_message = (
            response.text
        )

    raise RuntimeError(
        f"KiosAPI Error "
        f"{response.status_code}: "
        f"{error_message}"
    )

data = response.json()

try:

    content = (
        data["choices"][0]
        ["message"]["content"]
    )

except Exception:

    raise RuntimeError(
        "استجابة KiosAPI غير متوقعة."
    )

return normalize_results(
    parse_json(content)
)
```

# ============================================================

# OpenAI الرسمي

# ============================================================

def call_openai(
text,
api_key,
model
):

```
import requests

response = requests.post(

    "https://api.openai.com/v1/chat/completions",

    headers={
        "Authorization":
            f"Bearer {api_key}",

        "Content-Type":
            "application/json"
    },

    json={
        "model": model,

        "messages": [
            {
                "role": "system",
                "content":
                    "أنت محكّم أكاديمي. "
                    "أجب بالعربية وأعد JSON فقط."
            },
            {
                "role": "user",
                "content":
                    build_prompt(text)
            }
        ],

        "temperature": 0.1
    },

    timeout=300
)

if response.status_code != 200:

    raise RuntimeError(
        f"OpenAI Error "
        f"{response.status_code}: "
        f"{response.text}"
    )

content = (
    response.json()
    ["choices"][0]
    ["message"]
    ["content"]
)

return normalize_results(
    parse_json(content)
)
```

# ============================================================

# Gemini

# ============================================================

def call_gemini(
text,
api_key,
model
):

```
import requests

url = (
    "https://generativelanguage.googleapis.com/"
    f"v1beta/models/{model}:generateContent"
    f"?key={api_key}"
)

response = requests.post(

    url,

    headers={
        "Content-Type":
            "application/json"
    },

    json={
        "systemInstruction": {
            "parts": [
                {
                    "text":
                        "أنت محكّم أكاديمي. "
                        "أجب بالعربية وأعد JSON فقط."
                }
            ]
        },

        "contents": [
            {
                "role": "user",
                "parts": [
                    {
                        "text":
                            build_prompt(text)
                    }
                ]
            }
        ],

        "generationConfig": {
            "temperature": 0.1,
            "responseMimeType":
                "application/json"
        }
    },

    timeout=300
)

if response.status_code != 200:

    raise RuntimeError(
        f"Gemini Error "
        f"{response.status_code}: "
        f"{response.text}"
    )

content = (
    response.json()
    ["candidates"][0]
    ["content"]["parts"][0]
    ["text"]
)

return normalize_results(
    parse_json(content)
)
```

# ============================================================

# Claude عبر Anthropic الرسمي

# ============================================================

def call_anthropic(
text,
api_key,
model
):

```
import requests

response = requests.post(

    "https://api.anthropic.com/v1/messages",

    headers={
        "x-api-key":
            api_key,

        "anthropic-version":
            "2023-06-01",

        "content-type":
            "application/json"
    },

    json={
        "model": model,

        "max_tokens": 16000,

        "messages": [
            {
                "role": "user",
                "content":
                    build_prompt(text)
            }
        ]
    },

    timeout=300
)

if response.status_code != 200:

    raise RuntimeError(
        f"Anthropic Error "
        f"{response.status_code}: "
        f"{response.text}"
    )

content = "\n".join(

    block.get(
        "text",
        ""
    )

    for block in (
        response.json()
        .get(
            "content",
            []
        )
    )

    if block.get(
        "type"
    ) == "text"
)

return normalize_results(
    parse_json(content)
)
```

# ============================================================

# تقييم

# ============================================================

def evaluate(
text,
provider,
api_key,
model
):

```
if not api_key:

    raise ValueError(
        "لم يتم إدخال مفتاح API."
    )

if provider == "KiosAPI":

    return call_kiosapi(
        text,
        api_key,
        model
    )

if provider == "OpenAI":

    return call_openai(
        text,
        api_key,
        model
    )

if provider == "Gemini":

    return call_gemini(
        text,
        api_key,
        model
    )

if provider == "Claude الرسمي":

    return call_anthropic(
        text,
        api_key,
        model
    )

raise ValueError(
    "مزود API غير معروف."
)
```

# ============================================================

# حساب النتيجة

# ============================================================

def calculate(
results
):

```
total = sum(
    item["score"]
    for item in results
)

percentage = (
    total
    / MAX_SCORE
    * 100
    if MAX_SCORE
    else 0
)

if percentage >= 85:

    level = "ممتاز"

elif percentage >= 70:

    level = "جيد جداً"

elif percentage >= 55:

    level = "جيد"

elif percentage >= 40:

    level = "مقبول"

else:

    level = "يحتاج إلى تحسين"

return (
    total,
    MAX_SCORE,
    percentage,
    level
)
```

# ============================================================

# التقرير Markdown

# ============================================================

def build_markdown_report(
filename,
provider,
results,
analysis
):

```
total, maximum, percentage, level = (
    calculate(
        results
    )
)

lines = [

    "# تقرير تقييم خطة البحث العلمي",

    "",

    f"**الملف:** {filename}",

    f"**مزود الذكاء الاصطناعي:** {provider}",

    f"**التاريخ:** "
    f"{datetime.now().strftime('%Y-%m-%d %H:%M')}",

    "",

    "**نطاق التقييم:** "
    "المعايير 12–51 و68–71 فقط.",

    "",

    f"## النتيجة العامة",

    "",

    f"**{percentage:.1f}% — "
    f"{total}/{maximum} — "
    f"{level}**",

    "",

    "## التحليل الأكاديمي",

    "",

    analysis,

    "",

    "## التقييم التفصيلي",

    ""
]

current_category = None

for item in results:

    if (
        item["category"]
        != current_category
    ):

        lines.extend(
            [
                f"### {item['category']}",
                ""
            ]
        )

        current_category = (
            item["category"]
        )

    lines.extend(
        [
            f"#### المعيار {item['id']}",
            "",
            f"**المعيار:** "
            f"{item['criterion']}",
            "",
            f"**الدرجة:** "
            f"{item['score']}/2",
            "",
            f"**الحكم:** "
            f"{item['status']}",
            "",
            f"**الدليل:** "
            f"{item['evidence']}",
            "",
            f"**التحليل:** "
            f"{item['explanation']}",
            "",
            f"**التوصية:** "
            f"{item['suggestion']}",
            ""
        ]
    )

return "\n".join(
    lines
)
```

# ============================================================

# القائمة الجانبية

# ============================================================

with st.sidebar:

```
st.header(
    "⚙️ إعدادات التقييم"
)

mode = st.radio(
    "طريقة التقييم",

    [
        "🤖 تقييم بالذكاء الاصطناعي",
        "🔎 فحص أولي"
    ]
)

st.divider()

st.markdown(
    "### 📌 نطاق التقييم"
)

st.info(
    "المعايير 12–51 و68–71 فقط"
)

st.write(
    f"عدد المعايير: "
    f"**{len(CRITERIA)}**"
)

st.write(
    f"الدرجة القصوى: "
    f"**{MAX_SCORE}**"
)
```

# ============================================================

# إعداد مزود API

# ============================================================

if mode == "🤖 تقييم بالذكاء الاصطناعي":

```
provider = st.selectbox(
    "اختر مزود الذكاء الاصطناعي",

    [
        "KiosAPI",
        "OpenAI",
        "Gemini",
        "Claude الرسمي"
    ]
)

# --------------------------------------------------------
# KiosAPI
# --------------------------------------------------------

if provider == "KiosAPI":

    api_key = st.text_input(
        "KiosAPI API Key",

        value=get_secret(
            "KIOSAPI_API_KEY"
        ),

        type="password"
    )

    model = st.selectbox(
        "نموذج الذكاء الاصطناعي",

        [
            "anthropic/claude-opus-4-8",
            "anthropic/claude-sonnet-4-6"
        ]
    )

# --------------------------------------------------------
# OpenAI
# --------------------------------------------------------

elif provider == "OpenAI":

    api_key = st.text_input(
        "OpenAI API Key",

        value=get_secret(
            "OPENAI_API_KEY"
        ),

        type="password"
    )

    model = st.selectbox(
        "النموذج",

        [
            "gpt-4.1-mini",
            "gpt-4.1",
            "gpt-5"
        ]
    )

# --------------------------------------------------------
# Gemini
# --------------------------------------------------------

elif provider == "Gemini":

    api_key = st.text_input(
        "Gemini API Key",

        value=get_secret(
            "GEMINI_API_KEY"
        ),

        type="password"
    )

    model = st.selectbox(
        "النموذج",

        [
            "gemini-2.5-flash",
            "gemini-2.5-pro"
        ]
    )

# --------------------------------------------------------
# Claude الرسمي
# --------------------------------------------------------

else:

    api_key = st.text_input(
        "Anthropic API Key",

        value=get_secret(
            "ANTHROPIC_API_KEY"
        ),

        type="password"
    )

    model = st.selectbox(
        "نموذج Claude",

        [
            "claude-opus-4-8",
            "claude-sonnet-4-5"
        ]
    )
```

else:

```
provider = "فحص أولي"

api_key = ""

model = ""
```

# ============================================================

# رفع الملف

# ============================================================

uploaded_file = st.file_uploader(

```
"📤 ارفع خطة البحث",

type=[
    "pdf",
    "docx",
    "png",
    "jpg",
    "jpeg",
    "webp"
]
```

)

# ============================================================

# معالجة الملف

# ============================================================

if uploaded_file:

```
with st.spinner(
    "🔎 جاري استخراج النص..."
):

    try:

        extracted_text, extraction_method = (
            extract_file(
                uploaded_file
            )
        )

    except Exception as error:

        st.error(
            "حدث خطأ أثناء استخراج النص:"
        )

        st.exception(
            error
        )

        st.stop()

if not extracted_text:

    st.error(
        "لم يتم استخراج أي نص من الملف."
    )

    st.stop()

st.success(
    f"✅ تم استخراج النص بنجاح — "
    f"{extraction_method} — "
    f"{len(extracted_text):,} حرف."
)

# --------------------------------------------------------
# معاينة النص
# --------------------------------------------------------

with st.expander(
    "👁️ معاينة النص المستخرج",
    expanded=True
):

    st.text_area(
        "النص المستخرج",

        extracted_text,

        height=350,

        label_visibility="collapsed"
    )

# --------------------------------------------------------
# زر التقييم
# --------------------------------------------------------

if st.button(
    "🚀 ابدأ تقييم خطة البحث",

    type="primary",

    use_container_width=True
):

    try:

        if mode == "🤖 تقييم بالذكاء الاصطناعي":

            with st.spinner(
                f"🤖 جاري التقييم باستخدام "
                f"{provider}..."
            ):

                results, analysis = (
                    evaluate(
                        extracted_text,
                        provider,
                        api_key,
                        model
                    )
                )

        else:

            # ==========================================
            # فحص أولي
            # ==========================================

            results = []

            text_lower = (
                extracted_text
                .lower()
            )

            keywords = [
                "عنوان",
                "مقدمة",
                "مشكلة",
                "سؤال",
                "تساؤل",
                "فرض",
                "أهمية",
                "هدف",
                "مصطلح",
                "حدود",
                "منهج",
                "دراسة",
                "بحث",
                "لغة"
            ]

            hits = sum(
                1
                for keyword in keywords
                if keyword in text_lower
            )

            score = (
                2
                if hits >= 10
                else 1
                if hits >= 4
                else 0
            )

            for (
                number,
                category,
                criterion
            ) in CRITERIA:

                results.append(
                    {
                        "id":
                            number,

                        "category":
                            category,

                        "criterion":
                            criterion,

                        "score":
                            score,

                        "status":
                            [
                                "غير متحقق",
                                "متحقق جزئياً",
                                "متحقق"
                            ][score],

                        "evidence":
                            "نتيجة فحص أولي آلي.",

                        "explanation":
                            "هذه نتيجة أولية "
                            "تعتمد على مؤشرات لغوية "
                            "ولا تمثل حكماً أكاديمياً نهائياً.",

                        "suggestion":
                            "استخدم التقييم بالذكاء "
                            "الاصطناعي للحصول على "
                            "تحليل أكاديمي أدق."
                    }
                )

            analysis = (
                "هذه نتيجة فحص أولي آلي. "
                "لا ينبغي اعتمادها كتقييم أكاديمي نهائي."
            )

        # حفظ النتائج
        st.session_state[
            "results"
        ] = results

        st.session_state[
            "analysis"
        ] = analysis

        st.session_state[
            "filename"
        ] = uploaded_file.name

        st.session_state[
            "provider"
        ] = provider

        st.success(
            "✅ اكتمل تقييم خطة البحث."
        )

    except Exception as error:

        st.error(
            f"حدث خطأ أثناء التقييم: "
            f"{error}"
        )

        with st.expander(
            "تفاصيل الخطأ"
        ):

            st.exception(
                error
            )
```

# ============================================================

# عرض النتائج

# ============================================================

if "results" in st.session_state:

```
results = (
    st.session_state[
        "results"
    ]
)

analysis = (
    st.session_state[
        "analysis"
    ]
)

total, maximum, percentage, level = (
    calculate(
        results
    )
)

st.divider()

# --------------------------------------------------------
# المؤشرات
# --------------------------------------------------------

col1, col2, col3, col4 = (
    st.columns(4)
)

col1.metric(
    "النتيجة",
    f"{percentage:.1f}%"
)

col2.metric(
    "النقاط",
    f"{total}/{maximum}"
)

col3.metric(
    "التقدير",
    level
)

col4.metric(
    "عدد المعايير",
    len(results)
)

st.progress(
    percentage / 100
)

# --------------------------------------------------------
# التحليل العام
# --------------------------------------------------------

st.subheader(
    "📝 التحليل الأكاديمي العام"
)

st.markdown(
    f"""
    <div class="result-card">
    {analysis.replace(chr(10), "<br>")}
    </div>
    """,
    unsafe_allow_html=True
)

# --------------------------------------------------------
# ملخص المحاور
# --------------------------------------------------------

st.subheader(
    "📊 النتائج حسب المحاور"
)

summary = []

for (
    category,
    category_items
) in CATEGORIES.items():

    category_results = [
        item
        for item in results
        if item["category"]
        == category
    ]

    category_score = sum(
        item["score"]
        for item in category_results
    )

    category_max = (
        len(category_items)
        * 2
    )

    category_percentage = (
        category_score
        / category_max
        * 100
        if category_max
        else 0
    )

    summary.append(
        {
            "المحور":
                category,

            "عدد المعايير":
                len(category_items),

            "النقاط":
                f"{category_score}/"
                f"{category_max}",

            "النسبة":
                f"{category_percentage:.1f}%"
        }
    )

st.dataframe(
    summary,

    use_container_width=True,

    hide_index=True
)

# --------------------------------------------------------
# التقييم التفصيلي
# --------------------------------------------------------

st.subheader(
    "🧾 التقييم التفصيلي"
)

for category in CATEGORIES:

    st.markdown(
        f"### {category}"
    )

    category_results = [
        item
        for item in results
        if item["category"]
        == category
    ]

    for item in category_results:

        if item["score"] == 0:

            icon = "🔴"

        elif item["score"] == 1:

            icon = "🟠"

        else:

            icon = "🟢"

        with st.expander(
            f"{icon} المعيار "
            f"{item['id']} — "
            f"{item['criterion']} "
            f"— {item['score']}/2"
        ):

            st.write(
                f"**الحكم:** "
                f"{item['status']}"
            )

            st.markdown(
                "**📌 الدليل من الخطة:**"
            )

            st.info(
                item["evidence"]
            )

            st.markdown(
                "**🔎 التحليل الأكاديمي:**"
            )

            st.write(
                item["explanation"]
            )

            st.markdown(
                "**💡 التوصية:**"
            )

            if item["suggestion"]:

                st.warning(
                    item["suggestion"]
                )

            else:

                st.success(
                    "لا توجد توصية إضافية."
                )

# --------------------------------------------------------
# المعايير التي تحتاج إلى تحسين
# --------------------------------------------------------

st.subheader(
    "🎯 المعايير التي تحتاج إلى تحسين"
)

weak = [
    item
    for item in results
    if item["score"] < 2
]

if weak:

    for item in weak:

        st.markdown(
            f"**{item['id']}. "
            f"{item['criterion']}** — "
            f"{item['suggestion']}"
        )

else:

    st.success(
        "🎉 جميع المعايير حصلت على الدرجة الكاملة."
    )

# --------------------------------------------------------
# التقرير
# --------------------------------------------------------

markdown_report = (
    build_markdown_report(
        st.session_state[
            "filename"
        ],

        st.session_state[
            "provider"
        ],

        results,

        analysis
    )
)

json_report = {

    "filename":
        st.session_state[
            "filename"
        ],

    "provider":
        st.session_state[
            "provider"
        ],

    "evaluation_scope":
        "12–51 و68–71 فقط",

    "criteria_count":
        len(results),

    "total_score":
        total,

    "maximum_score":
        maximum,

    "percentage":
        percentage,

    "level":
        level,

    "overall_analysis":
        analysis,

    "results":
        results
}

st.subheader(
    "📥 حفظ التقرير"
)

col1, col2 = (
    st.columns(2)
)

with col1:

    st.download_button(

        "⬇️ تنزيل JSON",

        data=json.dumps(
            json_report,
            ensure_ascii=False,
            indent=2
        ),

        file_name=
            "تقرير_تقييم_خطة_البحث.json",

        mime=
            "application/json",

        use_container_width=True
    )

with col2:

    st.download_button(

        "⬇️ تنزيل Markdown",

        data=markdown_report,

        file_name=
            "تقرير_تقييم_خطة_البحث.md",

        mime=
            "text/markdown",

        use_container_width=True
    )
```

else:

```
st.info(
    "📤 ارفع خطة البحث للبدء في التقييم."
)
```
